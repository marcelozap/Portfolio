from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


OUT_DIR = Path(r"C:\Users\Green Machine\Desktop\Portfolio\career-kit")
DOCX_PATH = OUT_DIR / "Marcelo-Zapata-Foundation-Health-Software-Engineer-Resume.docx"
TXT_PATH = OUT_DIR / "Marcelo-Zapata-Foundation-Health-Software-Engineer-Resume.txt"
PDF_PATH = OUT_DIR / "Marcelo-Zapata-Foundation-Health-Software-Engineer-Resume.pdf"


content = {
    "name": "MARCELO ZAPATA",
    "contact": "Orlando, FL | marcelozapata08@gmail.com | linkedin.com/in/marcelozap | marcelozapata.dev",
    "summary": (
        "Software engineer with 4 years of experience across enterprise modernization, QA automation, "
        "internal workflow tooling, data/reporting workflows, and independent product builds. Strong fit "
        "for AI-enabled healthcare and pharmacy infrastructure: combines software engineering discipline, "
        "pharmacy operations context, AI-native development habits, and a builder mindset across C#, "
        "TypeScript, Python, SQL, Azure-based delivery workflows, Playwright, and API integrations."
    ),
    "skills": [
        ("Languages", "C#, TypeScript, Python, SQL, Swift"),
        ("Backend / Product", "APIs, FastAPI, .NET/Avalonia, React, Next.js, Power Platform, internal workflow tooling"),
        ("Cloud / Data / Delivery", "Azure, CI/CD, Snowflake, Databricks, Power BI, SQLite, secure configuration"),
        ("Testing / Automation", "Playwright, regression testing, QA automation, test planning, acceptance criteria"),
        ("AI / Domain", "AI-assisted development with ChatGPT/Codex/Cursor/Claude, healthcare/pharmacy workflow context"),
    ],
    "experience": [
        {
            "company": "Publix Super Markets",
            "role": "Software Engineer",
            "dates": "2022 - Present",
            "bullets": [
                "Led QA automation efforts for modernization work on a C#-heavy enterprise stack, building Playwright regression coverage and helping move validation into repeatable CI/CD workflows.",
                "Improved release confidence by translating legacy behavior, edge cases, and integration expectations into clearer test coverage, acceptance criteria, and implementation notes.",
                "Contributed to data, reporting, and workflow-tooling projects using Snowflake, Databricks, Power BI, Power Platform, and API integrations.",
                "Worked with secure test configuration and Azure-based delivery patterns, including pipeline execution and secret/configuration handling.",
                "Built and supported internal workflow tooling that reduced manual reconciliation, file-exception handling, and repetitive follow-up work.",
                "Documented workflows, test plans, system behavior, and handoff notes so teammates could understand changes, reproduce issues, and support releases more confidently.",
                "Use AI-assisted development tools to accelerate research, implementation planning, test design, documentation, and debugging while keeping engineering review discipline.",
            ],
        },
        {
            "company": "Florida State University",
            "role": "S.T.E.M. Course Instructor",
            "dates": "2019 - 2022",
            "bullets": [
                "Taught STEM-focused coursework and supported students through technical concepts, structured problem solving, and hands-on learning.",
                "Built communication habits that carry into engineering work: clear documentation, patient collaboration, and practical mentorship.",
            ],
        },
        {
            "company": "Pharmacy Operations Experience",
            "role": "Prior experience",
            "dates": "",
            "bullets": [
                "Bring firsthand context from pharmacy operations and customer/patient-facing workflows, including the importance of accuracy, coordination, and reducing friction in service delivery.",
            ],
        },
    ],
    "projects": [
        {
            "name": "GateKPT MusicOS",
            "dates": "2025 - Present",
            "desc": "C#/.NET Avalonia creative operating system for live-loop artists, combining performance planning, local project memory, lyric/caption workflows, hardware routing, and audio-reactive visual artwork.",
        },
        {
            "name": "Green Machine Quant OS",
            "dates": "2021 - Present",
            "desc": "Python/FastAPI research and paper-execution platform connecting OAuth-secured broker data with local backtesting, risk analytics, paper trading, and AI-assisted research review.",
        },
        {
            "name": "Rally",
            "dates": "2025 - Present",
            "desc": "Local-first iOS product using SwiftUI, SpriteKit, SwiftData, and AVAudioEngine for tennis training, journaling, court discovery, and daily engagement loops.",
        },
    ],
    "education": "Florida State University - Bachelor of Science in Computer Science & Business | 2019 - 2022",
}


def set_run(run, size=9.3, bold=False, color="000000"):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text.upper())
    set_run(r, size=9.5, bold=True, color="1F4D78")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(1.2)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_run(r, size=8.45)


def build_docx():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(9)
    styles["Normal"].paragraph_format.space_after = Pt(2)
    styles["Normal"].paragraph_format.line_spacing = 1.0

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    r = title.add_run(content["name"])
    set_run(r, size=17, bold=True, color="0B2545")

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(5)
    r = contact.add_run(content["contact"])
    set_run(r, size=8.7, color="333333")

    add_section_heading(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(content["summary"])
    set_run(r, size=8.8)

    add_section_heading(doc, "Technical Skills")
    for label, value in content["skills"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(f"{label}: ")
        set_run(r, size=8.55, bold=True)
        r = p.add_run(value)
        set_run(r, size=8.55)

    add_section_heading(doc, "Experience")
    for item in content["experience"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        left = p.add_run(item["company"])
        set_run(left, size=9.2, bold=True)
        if item["dates"]:
            right = p.add_run(f" | {item['dates']}")
            set_run(right, size=8.8, color="555555")

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(item["role"])
        set_run(r, size=8.8, bold=True, color="333333")

        for bullet in item["bullets"]:
            add_bullet(doc, bullet)

    add_section_heading(doc, "Selected Builds")
    for project in content["projects"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(project["name"])
        set_run(r, size=9, bold=True)
        r = p.add_run(f" | {project['dates']}")
        set_run(r, size=8.6, color="555555")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(project["desc"])
        set_run(r, size=8.4)

    add_section_heading(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run(content["education"])
    set_run(r, size=8.7)

    doc.save(DOCX_PATH)


def build_txt():
    lines = [
        content["name"],
        content["contact"],
        "",
        "PROFESSIONAL SUMMARY",
        content["summary"],
        "",
        "TECHNICAL SKILLS",
    ]
    for label, value in content["skills"]:
        lines.append(f"{label}: {value}")
    lines.extend(["", "EXPERIENCE"])
    for item in content["experience"]:
        heading = f"{item['company']} | {item['role']}"
        if item["dates"]:
            heading += f" | {item['dates']}"
        lines.append(heading)
        for bullet in item["bullets"]:
            lines.append(f"- {bullet}")
        lines.append("")
    lines.append("SELECTED BUILDS")
    for project in content["projects"]:
        lines.append(f"{project['name']} | {project['dates']}")
        lines.append(f"- {project['desc']}")
        lines.append("")
    lines.extend(["EDUCATION", content["education"], ""])
    TXT_PATH.write_text("\n".join(lines), encoding="utf-8")


def build_pdf():
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=0.48 * inch,
        rightMargin=0.48 * inch,
        topMargin=0.42 * inch,
        bottomMargin=0.42 * inch,
    )
    styles = {
        "name": ParagraphStyle(
            "name",
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=17,
            alignment=1,
            textColor=colors.HexColor("#0B2545"),
            spaceAfter=1,
        ),
        "contact": ParagraphStyle(
            "contact",
            fontName="Helvetica",
            fontSize=8.5,
            leading=10,
            alignment=1,
            textColor=colors.HexColor("#333333"),
            spaceAfter=5,
        ),
        "section": ParagraphStyle(
            "section",
            fontName="Helvetica-Bold",
            fontSize=9.5,
            leading=10.5,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=5,
            spaceAfter=1,
        ),
        "body": ParagraphStyle(
            "body",
            fontName="Helvetica",
            fontSize=8.4,
            leading=9.5,
            spaceAfter=2,
        ),
        "small": ParagraphStyle(
            "small",
            fontName="Helvetica",
            fontSize=8.1,
            leading=9.1,
            spaceAfter=1.4,
        ),
        "bullet": ParagraphStyle(
            "bullet",
            fontName="Helvetica",
            fontSize=7.55,
            leading=8.45,
            leftIndent=10,
            firstLineIndent=-6,
            bulletIndent=0,
            spaceAfter=1.0,
        ),
        "role": ParagraphStyle(
            "role",
            fontName="Helvetica-Bold",
            fontSize=8.4,
            leading=9.2,
            spaceAfter=0.5,
        ),
    }

    story = [
        Paragraph(content["name"], styles["name"]),
        Paragraph(content["contact"], styles["contact"]),
        Paragraph("PROFESSIONAL SUMMARY", styles["section"]),
        Paragraph(content["summary"], styles["body"]),
        Paragraph("TECHNICAL SKILLS", styles["section"]),
    ]
    for label, value in content["skills"]:
        story.append(Paragraph(f"<b>{label}:</b> {value}", styles["small"]))

    story.append(Paragraph("EXPERIENCE", styles["section"]))
    for item in content["experience"]:
        title = item["company"]
        if item["dates"]:
            title += f" | <font color='#555555'>{item['dates']}</font>"
        story.append(Paragraph(f"<b>{title}</b>", styles["body"]))
        story.append(Paragraph(item["role"], styles["role"]))
        for bullet in item["bullets"]:
            story.append(Paragraph(bullet, styles["bullet"], bulletText="-"))

    story.append(Paragraph("SELECTED BUILDS", styles["section"]))
    for project in content["projects"]:
        story.append(
            Paragraph(
                f"<b>{project['name']}</b> | <font color='#555555'>{project['dates']}</font>",
                styles["body"],
            )
        )
        story.append(Paragraph(project["desc"], styles["small"]))

    story.append(Paragraph("EDUCATION", styles["section"]))
    story.append(Paragraph(content["education"], styles["small"]))
    doc.build(story)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_txt()
    build_pdf()
    print(DOCX_PATH)
    print(TXT_PATH)
    print(PDF_PATH)
